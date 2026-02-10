#!/usr/bin/env python3
"""
generate_design_doc.py
Generates the ADS-507 Team 4 Design Document in Word (.docx) format.
Run:  python3 docs/generate_design_doc.py
Output: docs/Final-Project-Report-Team-4.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Final-Project-Report-Team-4.docx")


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    shading.append(shading_elem)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E4057")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    return table


def build_document():
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # =====================================================================
    # TITLE PAGE
    # =====================================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ADS-507 Final Project\nDesign Document")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "E-Commerce Data Pipeline for Order and Delivery Insights"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

    doc.add_paragraph("")

    team_info = doc.add_paragraph()
    team_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team_info.add_run(
        "Team 4\n\n"
        "Nikita Rogers\n"
        "Jun Sik Ryu\n"
        "Faye Shawntel Corprew\n\n"
        "University of San Diego\n"
        "ADS-507 Practical Data Engineering\n"
        "Spring 2026"
    )
    run.font.size = Pt(12)

    doc.add_paragraph("")
    repo_link = doc.add_paragraph()
    repo_link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = repo_link.add_run(
        "GitHub Repository: https://github.com/nikitarogers333/ads507-team4-project"
    )
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

    doc.add_page_break()

    # =====================================================================
    # TABLE OF CONTENTS
    # =====================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Source Data",
        "3. System Architecture",
        "4. ETL Pipeline Design",
        "5. Database Schema",
        "6. Pipeline Output",
        "7. Monitoring and Validation",
        "8. Deployment and Reproducibility",
        "9. Addressing Instructor Feedback",
        "10. System Gaps and Next Steps",
        "11. References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # =====================================================================
    # 1. INTRODUCTION
    # =====================================================================
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document describes the design and implementation of a production-ready "
        "ETL (Extract-Transform-Load) data pipeline developed for the ADS-507 Practical "
        "Data Engineering course at the University of San Diego. The pipeline processes "
        "the Brazilian E-Commerce dataset published by Olist on Kaggle, transforming raw "
        "transactional data into a star schema optimized for analytical queries."
    )
    doc.add_paragraph(
        "The system is fully containerized using Docker Compose and can be deployed with "
        "a single command. It downloads raw CSV data from a GitHub release, loads it into "
        "MySQL 8.0 staging tables, applies SQL-based transformations to build dimension "
        "and fact tables, and produces five analytical views that provide insights into "
        "revenue trends, delivery performance, seller metrics, product category analysis, "
        "and customer segmentation."
    )
    doc.add_paragraph(
        "The GitHub repository containing all source code is available at: "
        "https://github.com/nikitarogers333/ads507-team4-project"
    )

    # =====================================================================
    # 2. SOURCE DATA
    # =====================================================================
    doc.add_heading("2. Source Data", level=1)

    doc.add_heading("2.1 Dataset Selection", level=2)
    doc.add_paragraph(
        "We selected the Brazilian E-Commerce Public Dataset by Olist, available on Kaggle "
        "(https://www.kaggle.com/datasets/olistbr/brazilian-ecommerce). This dataset was "
        "chosen because:"
    )
    bullets = [
        "It contains multiple related tables (9 CSV files) that naturally support relational "
        "modeling, making it ideal for practicing SQL-based transformations.",
        "With approximately 100,000 orders and over 1 million geolocation records, the dataset "
        "is large enough to demonstrate real-world pipeline challenges without requiring "
        "specialized infrastructure.",
        "The data covers the period from 2016 to 2018 and includes transactional, categorical, "
        "temporal, and geographic data types, providing rich opportunities for analytical queries.",
        "Multiple entities (customers, orders, items, payments, reviews, products, sellers) "
        "create natural join relationships that support a star schema design.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2.2 Dataset Description", level=2)
    add_table(
        doc,
        ["CSV File", "Description", "Approx. Rows", "Key Columns"],
        [
            ["olist_customers_dataset.csv", "Customer profiles", "99,441", "customer_id, city, state"],
            ["olist_orders_dataset.csv", "Order headers with timestamps", "99,441", "order_id, status, dates"],
            ["olist_order_items_dataset.csv", "Line items per order", "112,650", "product_id, seller_id, price"],
            ["olist_order_payments_dataset.csv", "Payment records", "103,886", "payment_type, value"],
            ["olist_order_reviews_dataset.csv", "Customer reviews", "99,224", "review_score, comments"],
            ["olist_products_dataset.csv", "Product catalog", "32,951", "category, dimensions, weight"],
            ["olist_sellers_dataset.csv", "Seller profiles", "3,095", "seller_id, city, state"],
            ["olist_geolocation_dataset.csv", "Zip code coordinates", "~1,000,000", "zip, lat, lng"],
            ["product_category_name_translation.csv", "Portuguese to English", "71", "category names"],
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "The geolocation dataset was split into 11 parts for the GitHub release due to file "
        "size constraints. The pipeline automatically reassembles these parts during the "
        "extraction phase."
    )

    # =====================================================================
    # 3. SYSTEM ARCHITECTURE
    # =====================================================================
    doc.add_heading("3. System Architecture", level=1)

    doc.add_heading("3.1 Architecture Overview", level=2)
    doc.add_paragraph(
        "The system uses a containerized architecture orchestrated by Docker Compose. "
        "Three services collaborate to form the complete pipeline:"
    )
    add_table(
        doc,
        ["Service", "Image", "Purpose", "Port"],
        [
            ["mysql", "mysql:8.0", "Persistent data store (staging, dim, fact, views)", "3306"],
            ["pipeline", "alpine:3.19", "ETL orchestrator (download, load, transform, validate)", "—"],
            ["adminer", "adminer:4", "Web-based database GUI for monitoring and queries", "8081"],
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "The architecture follows the Infrastructure as Code (IaC) principle: the entire "
        "system is defined in docker-compose.yml and can be deployed with "
        "'docker compose up'. No manual infrastructure setup is required."
    )

    doc.add_heading("3.2 Architecture Diagram", level=2)
    doc.add_paragraph(
        "The following diagram illustrates the data flow through the system:"
    )
    # Text-based architecture diagram
    arch_diagram = (
        "┌─────────────────────────────────────────────────────────────┐\n"
        "│                Docker Compose Environment                   │\n"
        "│                                                             │\n"
        "│  ┌──────────┐    ┌─────────────────────────┐  ┌─────────┐  │\n"
        "│  │  GitHub   │───▶│   Pipeline Container    │  │ Adminer │  │\n"
        "│  │  Release  │    │   (Alpine Linux)        │  │ Web GUI │  │\n"
        "│  │ (9 CSVs)  │    │                         │  │  :8081  │  │\n"
        "│  └──────────┘    │  1. Download CSVs        │  └────┬────┘  │\n"
        "│                   │  2. Load → staging       │       │      │\n"
        "│                   │  3. Clean data           │       │      │\n"
        "│                   │  4. Build dimensions     │       │      │\n"
        "│                   │  5. Build facts          │       │      │\n"
        "│                   │  6. Create views         │       │      │\n"
        "│                   │  7. Validate             │       │      │\n"
        "│                   └────────────┬────────────┘       │      │\n"
        "│                                │                     │      │\n"
        "│                                ▼                     │      │\n"
        "│                   ┌─────────────────────────┐       │      │\n"
        "│                   │    MySQL 8.0 (:3306)    │◀──────┘      │\n"
        "│                   │  Staging │ Dimensions   │              │\n"
        "│                   │  Facts   │ Views        │              │\n"
        "│                   │  (persistent volume)    │              │\n"
        "│                   └─────────────────────────┘              │\n"
        "└─────────────────────────────────────────────────────────────┘"
    )
    p = doc.add_paragraph()
    run = p.add_run(arch_diagram)
    run.font.name = "Courier New"
    run.font.size = Pt(7.5)

    # =====================================================================
    # 4. ETL PIPELINE DESIGN
    # =====================================================================
    doc.add_heading("4. ETL Pipeline Design", level=1)

    doc.add_heading("4.1 Extract Phase", level=2)
    doc.add_paragraph(
        "The extraction phase downloads raw CSV files from a GitHub release "
        "(tag: v1.0-raw-data). The pipeline script uses curl to download each file "
        "and stores them in a shared Docker volume (/data/raw). The geolocation dataset, "
        "which was split into 11 parts due to GitHub's file size limits, is automatically "
        "reassembled by concatenating the parts with proper header handling. All files are "
        "then converted from Windows (CRLF) to Unix (LF) line endings using dos2unix."
    )

    doc.add_heading("4.2 Load Phase", level=2)
    doc.add_paragraph(
        "The load phase uses MySQL's LOAD DATA INFILE command to bulk-load CSV data into "
        "9 staging tables. This approach was chosen over row-by-row INSERT for performance: "
        "LOAD DATA INFILE can load hundreds of thousands of rows in seconds. The MySQL "
        "server is configured with --secure-file-priv=/data/raw to restrict file access "
        "to the designated data directory, and --local-infile=1 to enable the feature."
    )

    doc.add_heading("4.3 Transform Phase", level=2)
    doc.add_paragraph(
        "The transformation phase executes four SQL scripts in sequence:"
    )
    add_table(
        doc,
        ["Script", "Purpose", "Key Operations"],
        [
            [
                "010_clean_staging.sql",
                "Data cleaning",
                "Trim whitespace, normalise empty strings to NULL, "
                "standardise state codes to uppercase, lowercase category names",
            ],
            [
                "020_dim_tables.sql",
                "Build dimension tables",
                "Create dim_customers, dim_products (with English category names via JOIN), "
                "dim_sellers, dim_date (recursive CTE), dim_geography (aggregated coordinates)",
            ],
            [
                "030_fact_tables.sql",
                "Build fact tables",
                "Create fact_orders (enriched with totals, delivery metrics, late flag), "
                "fact_order_items (linked to dimension keys), fact_payments, fact_reviews",
            ],
            [
                "040_analytical_views.sql",
                "Create analytical views",
                "5 views: monthly revenue, delivery performance, seller scoreboard, "
                "product categories, customer segments (RFM-style)",
            ],
        ],
    )

    doc.add_heading("4.4 Validate Phase", level=2)
    doc.add_paragraph(
        "After transformations complete, the pipeline runs a comprehensive validation "
        "script (050_validate.sql) that checks:"
    )
    checks = [
        "Row counts for all staging, dimension, and fact tables",
        "Null checks on critical foreign key columns",
        "Referential integrity between fact and dimension tables",
        "Business rules (no negative payments, review scores within 1–5 range)",
        "Sample output from all analytical views",
    ]
    for c in checks:
        doc.add_paragraph(c, style="List Bullet")

    # =====================================================================
    # 5. DATABASE SCHEMA
    # =====================================================================
    doc.add_heading("5. Database Schema", level=1)

    doc.add_heading("5.1 Star Schema Design", level=2)
    doc.add_paragraph(
        "The transformed data follows a star schema design with 5 dimension tables "
        "and 4 fact tables. This design optimizes for analytical queries by denormalising "
        "data into a structure that minimises joins while maintaining data integrity."
    )

    # Schema diagram
    schema_text = (
        "                    ┌──────────────┐\n"
        "                    │  dim_date     │\n"
        "                    │  (date_key PK)│\n"
        "                    └──────┬───────┘\n"
        "                           │\n"
        "┌──────────────┐  ┌───────┴────────┐  ┌──────────────┐\n"
        "│dim_customers │──│  fact_orders    │──│ dim_geography │\n"
        "│(customer_key)│  │  (order_key PK) │  │ (geo_key PK) │\n"
        "└──────────────┘  └───────┬────────┘  └──────────────┘\n"
        "                          │\n"
        "                  ┌───────┴────────┐\n"
        "                  │fact_order_items │\n"
        "                  │ (item_key PK)  │\n"
        "                  └──┬──────────┬──┘\n"
        "                     │          │\n"
        "            ┌────────┴──┐  ┌────┴─────────┐\n"
        "            │dim_products│  │ dim_sellers   │\n"
        "            │(product_key│  │ (seller_key)  │\n"
        "            └───────────┘  └──────────────┘\n"
        "\n"
        "        ┌───────────────┐    ┌──────────────┐\n"
        "        │ fact_payments  │    │ fact_reviews  │\n"
        "        │(payment_key)  │    │(review_key)   │\n"
        "        └───────────────┘    └──────────────┘"
    )
    p = doc.add_paragraph()
    run = p.add_run(schema_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    doc.add_heading("5.2 Dimension Tables", level=2)
    add_table(
        doc,
        ["Table", "Primary Key", "Columns", "Source"],
        [
            ["dim_customers", "customer_key (auto)", "customer_id, unique_id, city, state, zip", "stg_customers"],
            ["dim_products", "product_key (auto)", "product_id, category (PT+EN), dimensions, weight", "stg_products + stg_category_translation"],
            ["dim_sellers", "seller_key (auto)", "seller_id, city, state, zip", "stg_sellers"],
            ["dim_date", "date_key (YYYYMMDD)", "full_date, year, quarter, month, day, day_name, week", "Generated (recursive CTE, 2016–2018)"],
            ["dim_geography", "geo_key (auto)", "zip_code_prefix, city, state, avg_lat, avg_lng", "stg_geolocation (aggregated)"],
        ],
    )

    doc.add_heading("5.3 Fact Tables", level=2)
    add_table(
        doc,
        ["Table", "Primary Key", "Foreign Keys", "Measures"],
        [
            ["fact_orders", "order_key", "customer_key, date_keys", "total_items, total_amount, total_freight, total_payment, delivery_days, is_late"],
            ["fact_order_items", "item_key", "product_key, seller_key", "price, freight_value"],
            ["fact_payments", "payment_key", "order_id", "payment_value, installments"],
            ["fact_reviews", "review_key", "order_id", "review_score, comments"],
        ],
    )

    # =====================================================================
    # 6. PIPELINE OUTPUT
    # =====================================================================
    doc.add_heading("6. Pipeline Output", level=1)

    doc.add_paragraph(
        "The pipeline produces five SQL views that provide actionable business intelligence. "
        "These views can be queried directly via Adminer (web GUI), any MySQL client, or "
        "connected to a BI tool like Tableau or Power BI."
    )

    doc.add_heading("6.1 Monthly Revenue (vw_monthly_revenue)", level=2)
    doc.add_paragraph(
        "Aggregates revenue, order volume, average order value, and total items sold by "
        "month. Useful for identifying seasonal trends and growth patterns in the e-commerce "
        "platform."
    )

    doc.add_heading("6.2 Delivery Performance (vw_delivery_performance)", level=2)
    doc.add_paragraph(
        "Breaks down delivery speed and late-delivery rates by customer state and month. "
        "Identifies regions with logistics challenges and helps prioritise delivery "
        "infrastructure improvements."
    )

    doc.add_heading("6.3 Seller Performance (vw_seller_performance)", level=2)
    doc.add_paragraph(
        "Ranks sellers by total revenue, order volume, average review score, and freight "
        "costs. Enables marketplace operators to identify top performers and underperformers."
    )

    doc.add_heading("6.4 Product Category Performance (vw_product_category_performance)", level=2)
    doc.add_paragraph(
        "Analyses sales volume, revenue, and customer satisfaction by product category "
        "(in English). Supports merchandising and inventory planning decisions."
    )

    doc.add_heading("6.5 Customer Segments (vw_customer_segments)", level=2)
    doc.add_paragraph(
        "Classifies customers into three segments based on purchase frequency: "
        "one-time, returning (2 orders), and loyal (3+ orders). Includes lifetime value, "
        "average order value, and customer tenure. Supports targeted marketing strategies."
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Why is this output useful? These views transform raw transactional data into "
        "decision-ready insights. A marketplace operator can use the delivery performance "
        "view to identify states with high late-delivery rates and allocate logistics "
        "resources accordingly. The seller performance view enables data-driven decisions "
        "about which sellers to promote or flag for quality improvement. The customer "
        "segmentation view supports retention strategies by identifying at-risk customers."
    )

    # =====================================================================
    # 7. MONITORING AND VALIDATION
    # =====================================================================
    doc.add_heading("7. Monitoring and Validation", level=1)

    doc.add_heading("7.1 Pipeline Monitoring", level=2)
    doc.add_paragraph("The system provides four monitoring mechanisms:")
    monitoring = [
        "Real-time logs: 'docker compose logs -f pipeline' streams pipeline output as it runs, "
        "showing each step's progress and timing.",
        "Monitoring dashboard: 'scripts/monitor.sh' displays a summary of database connection "
        "status, table sizes, pipeline completion status, and active MySQL processes.",
        "Adminer web GUI: Available at http://localhost:8081, provides a graphical interface "
        "to browse tables, run queries, and inspect data.",
        "Log files: Each pipeline run creates a timestamped log file in the pipeline_logs "
        "volume for audit and debugging purposes.",
    ]
    for m in monitoring:
        doc.add_paragraph(m, style="List Bullet")

    doc.add_heading("7.2 Data Validation", level=2)
    doc.add_paragraph(
        "The validation script (050_validate.sql) runs automatically at the end of each "
        "pipeline execution and can be run independently via 'scripts/validate.sh'. "
        "It performs row count verification, null checks on foreign keys, referential "
        "integrity validation, business rule enforcement, and samples from analytical views."
    )

    doc.add_heading("7.3 Automated Testing", level=2)
    doc.add_paragraph(
        "The test suite (tests/test_pipeline.sh) provides 27 integration tests covering "
        "container health, staging/dimension/fact table population, analytical view "
        "functionality, and data integrity rules. Tests can be run with 'make test'."
    )

    doc.add_heading("7.4 Continuous Integration", level=2)
    doc.add_paragraph(
        "GitHub Actions CI (.github/workflows/ci.yml) runs on every push and pull request "
        "to the main branch. It validates Docker Compose configuration, checks SQL syntax, "
        "lints shell scripts with ShellCheck, verifies all required files exist, and "
        "performs a Docker build test that starts MySQL and verifies staging table creation."
    )

    # =====================================================================
    # 8. DEPLOYMENT AND REPRODUCIBILITY
    # =====================================================================
    doc.add_heading("8. Deployment and Reproducibility", level=1)

    doc.add_paragraph(
        "The pipeline is fully reproducible from the GitHub repository. Deployment requires "
        "only Docker Desktop and Git:"
    )
    steps = [
        "Clone the repository: git clone https://github.com/nikitarogers333/ads507-team4-project.git",
        "Create environment file: cp .env.example .env",
        "Run the pipeline: docker compose up",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"Step {i}: {s}", style="List Number")

    doc.add_paragraph("")
    doc.add_paragraph(
        "The system follows the Infrastructure as Code (IaC) principle: all services, "
        "configurations, and SQL scripts are version-controlled in the repository. "
        "Docker Compose ensures consistent environments across all team members' machines, "
        "eliminating 'works on my machine' issues. The raw dataset is hosted as a GitHub "
        "release, ensuring it is always available and version-matched with the pipeline code."
    )

    # =====================================================================
    # 9. ADDRESSING INSTRUCTOR FEEDBACK
    # =====================================================================
    doc.add_heading("9. Addressing Instructor Feedback", level=1)

    doc.add_heading("9.1 Module 3 Proposal Feedback", level=2)
    doc.add_paragraph(
        'The instructor noted that our proposal was "well aligned with the goals of ADS-507" '
        "and that the Olist dataset was appropriate for a data engineering project. The "
        "feedback emphasized the importance of clearly defining the pipeline's transformation "
        "steps and output. We addressed this by:"
    )
    feedback_items = [
        "Implementing four clearly defined transformation scripts (clean, dimensions, facts, views) "
        "with detailed SQL comments explaining each operation.",
        "Creating five analytical views that produce concrete, useful business insights rather "
        "than generic aggregations.",
        "Documenting all SQL transformations in both the README and this design document.",
    ]
    for f in feedback_items:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("9.2 Module 5 Status Report", level=2)
    doc.add_paragraph(
        "At the time of the Module 5 submission, our database tables existed but were empty "
        "because data ingestion and SQL transformation steps were still being finalized. "
        "We identified the following challenges and how we resolved them:"
    )
    add_table(
        doc,
        ["Challenge (Module 5)", "Resolution (Module 7)"],
        [
            [
                "Database tables empty – data ingestion not complete",
                "Implemented automated data download from GitHub release and LOAD DATA INFILE "
                "for all 9 CSV files. Pipeline now loads ~500K+ rows automatically.",
            ],
            [
                "SQL transformations not finalized",
                "Created 4 transformation scripts with data cleaning, 5 dimension tables, "
                "4 fact tables, and 5 analytical views.",
            ],
            [
                "Data validation not consistent across environments",
                "Containerized entire pipeline with Docker Compose. All team members run "
                "identical environments. Added comprehensive validation script.",
            ],
            [
                "Need best practices for documenting SQL transformations",
                "Each SQL file includes detailed comments. README documents all transformations. "
                "Validation script verifies data quality after each run.",
            ],
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "No specific instructor feedback comments were received on the Module 5 submission. "
        "However, we proactively addressed all gaps identified in our own status report."
    )

    # =====================================================================
    # 10. SYSTEM GAPS AND NEXT STEPS
    # =====================================================================
    doc.add_heading("10. System Gaps and Next Steps", level=1)

    doc.add_heading("10.1 Scalability", level=2)
    doc.add_paragraph(
        "The current system processes approximately 100K orders in under 2 minutes on a "
        "standard laptop. However, as data volume grows to millions of records, several "
        "bottlenecks would emerge:"
    )
    scale_items = [
        "LOAD DATA INFILE performance degrades with very large files. A chunked loading "
        "strategy (loading data in batches) would improve performance and enable incremental updates.",
        "The single MySQL instance would need to be replaced with a distributed database "
        "or read replicas for concurrent analytical queries at scale.",
        "The geolocation table (~1M rows) could benefit from spatial indexing for "
        "geographic queries.",
    ]
    for s in scale_items:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("10.2 Security", level=2)
    doc.add_paragraph(
        "Current security gaps and recommended improvements:"
    )
    security_items = [
        "Database credentials are stored in a .env file. In production, these should be "
        "managed by a secrets manager (e.g., AWS Secrets Manager, HashiCorp Vault).",
        "The MySQL root account is used for pipeline operations. A least-privilege approach "
        "with separate accounts for loading, transforming, and reading would be more secure.",
        "Network access to MySQL (port 3306) is exposed on the host. In production, this "
        "should be restricted to internal Docker network only.",
        "The .env file is excluded from Git via .gitignore, preventing accidental "
        "credential exposure in the repository.",
    ]
    for s in security_items:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("10.3 Extensibility", level=2)
    doc.add_paragraph(
        "The system is designed for extensibility:"
    )
    ext_items = [
        "New data sources can be added by creating additional staging tables in sql/init/ "
        "and corresponding LOAD DATA statements in sql/load/.",
        "New transformations can be added as numbered SQL files in sql/transformations/ – "
        "the pipeline automatically executes all files matching the 0*.sql pattern.",
        "The analytical views can be extended or replaced without modifying the underlying "
        "fact and dimension tables.",
        "A time-based trigger (e.g., cron job or Apache Airflow DAG) could automate "
        "pipeline execution on a schedule.",
        "The pipeline output could be extended to include email alerts (e.g., when late "
        "delivery rate exceeds a threshold) or dashboard integration with Grafana.",
    ]
    for e in ext_items:
        doc.add_paragraph(e, style="List Bullet")

    doc.add_heading("10.4 Additional Future Improvements", level=2)
    improvements = [
        "Implement incremental loading (only process new/changed records) instead of "
        "full reload for each pipeline run.",
        "Add data lineage tracking to trace how each record was transformed.",
        "Implement a data catalog for self-service discovery of available tables and views.",
        "Add Grafana or Metabase as a containerized dashboard service for visual analytics.",
        "Implement CDC (Change Data Capture) for real-time data processing.",
    ]
    for imp in improvements:
        doc.add_paragraph(imp, style="List Bullet")

    # =====================================================================
    # 11. REFERENCES
    # =====================================================================
    doc.add_heading("11. References", level=1)
    references = [
        "Docker Inc. (2024). Docker Compose documentation. https://docs.docker.com/compose/",
        "Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit: The Definitive Guide "
        "to Dimensional Modeling (3rd ed.). Wiley.",
        "Olist. (2018). Brazilian E-Commerce Public Dataset by Olist [Dataset]. Kaggle. "
        "https://www.kaggle.com/datasets/olistbr/brazilian-ecommerce",
        "Oracle Corporation. (2024). MySQL 8.0 Reference Manual. "
        "https://dev.mysql.com/doc/refman/8.0/en/",
        "Oracle Corporation. (2024). LOAD DATA Statement. "
        "https://dev.mysql.com/doc/refman/8.0/en/load-data.html",
    ]
    for ref in references:
        doc.add_paragraph(ref, style="List Number")

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"Design document saved to: {OUTPUT_FILE}")
    print("Convert to PDF:  Open in Word/Google Docs → File → Download as PDF")


if __name__ == "__main__":
    build_document()
